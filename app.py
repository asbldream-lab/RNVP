"""
YouTube Transcript Extractor v2
--------------------------------
Colle l'URL d'une chaîne (ou d'une playlist), choisis le nombre de vidéos,
le script récupère les transcriptions et te génère un .docx.
"""

import os
import re
import html
import glob
import tempfile
from io import BytesIO

import streamlit as st
import yt_dlp
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound
from docx import Document
from docx.shared import Pt, RGBColor


# ============================================================================
# CONFIG STREAMLIT
# ============================================================================
st.set_page_config(
    page_title="YouTube Transcript Extractor",
    page_icon="🎬",
    layout="wide",
)

st.title("🎬 YouTube Transcript Extractor")
st.caption("Colle une URL de chaîne / playlist → récupère les N dernières transcriptions → télécharge un .docx")


# ============================================================================
# HELPERS - LISTING
# ============================================================================
def normalize_channel_url(url: str) -> str:
    """Force le suffixe /videos pour récupérer les uploads récents d'une chaîne."""
    url = url.strip().rstrip("/")
    if "playlist" in url or "watch?v=" in url or "/videos" in url:
        return url
    return url + "/videos"


def list_videos(url: str, n: int):
    """Liste les N dernières vidéos d'une chaîne ou playlist via yt-dlp."""
    ydl_opts = {
        "quiet": True,
        "no_warnings": True,
        "extract_flat": True,
        "playlistend": n,
        "skip_download": True,
    }
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=False)

    entries = info.get("entries") or []
    videos = []
    for e in entries[:n]:
        vid_id = e.get("id")
        if not vid_id:
            continue
        videos.append({
            "id": vid_id,
            "title": e.get("title", "Sans titre"),
            "url": e.get("url") or f"https://www.youtube.com/watch?v={vid_id}",
        })
    source_title = info.get("title") or "YouTube Transcripts"
    return videos, source_title


# ============================================================================
# HELPERS - LANG MATCHING
# ============================================================================
def match_lang(pref: str, pool: list):
    """
    Matche 'fr' avec 'fr', 'fr-FR', 'fr-CA'.
    Matche 'en' avec 'en', 'en-US', 'en-GB', 'en-orig'.
    Retourne le code exact présent dans le pool, ou None.
    """
    pref_lower = pref.lower()
    # 1. Match exact
    for lang in pool:
        if lang.lower() == pref_lower:
            return lang
    # 2. Match par préfixe avec tiret/underscore (fr-FR, en-US...)
    for lang in pool:
        lang_low = lang.lower()
        if lang_low.startswith(pref_lower + "-") or lang_low.startswith(pref_lower + "_"):
            return lang
    # 3. Match loose (base du code de langue)
    for lang in pool:
        base = lang.lower().split("-")[0].split("_")[0]
        if base == pref_lower:
            return lang
    return None


# ============================================================================
# HELPERS - VTT PARSING
# ============================================================================
def parse_vtt(path: str) -> str:
    """Extrait le texte propre d'un fichier .vtt (sans timestamps ni balises)."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception:
        return ""

    lines = []
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if "-->" in line:
            continue
        if line.startswith(("WEBVTT", "Kind:", "Language:", "NOTE", "STYLE", "REGION")):
            continue
        if re.match(r"^\d+$", line):
            continue
        # Retirer balises <c>, <00:00:00.000>, <v Speaker>, etc.
        line = re.sub(r"<[^>]+>", "", line)
        # Décoder &amp;, &#39;, etc.
        line = html.unescape(line)
        line = line.strip()
        if line and (not lines or lines[-1] != line):
            lines.append(line)
    return " ".join(lines)


# ============================================================================
# FETCH - MÉTHODE 1 : yt-dlp (primaire, plus fiable)
# ============================================================================
def fetch_transcript_ytdlp(video_id: str, lang_pref: list, include_auto: bool):
    """
    1. Liste les sous-titres disponibles via extract_info
    2. Détermine la meilleure langue cible
    3. Télécharge UNIQUEMENT cette langue en .vtt
    Retourne (texte, label_langue, message_erreur).
    """
    video_url = f"https://www.youtube.com/watch?v={video_id}"

    # Étape 1 : inventaire des sous-titres dispos
    try:
        with yt_dlp.YoutubeDL({"quiet": True, "no_warnings": True, "skip_download": True}) as ydl:
            info = ydl.extract_info(video_url, download=False)
    except Exception as e:
        return None, None, f"yt-dlp extract_info : {type(e).__name__}"

    manual_subs = list((info.get("subtitles") or {}).keys())
    auto_subs = list((info.get("automatic_captions") or {}).keys())

    if not manual_subs and not auto_subs:
        return None, None, "aucun sous-titre (ni manuel ni auto)"

    # Étape 2 : choisir la meilleure langue
    target_lang, target_type = None, None

    # Priorité 1 : manuels dans langues préférées
    for lang in lang_pref:
        m = match_lang(lang, manual_subs)
        if m:
            target_lang, target_type = m, "manuel"
            break

    # Priorité 2 : auto dans langues préférées
    if not target_lang and include_auto:
        for lang in lang_pref:
            m = match_lang(lang, auto_subs)
            if m:
                target_lang, target_type = m, "auto"
                break

    # Priorité 3 : n'importe quel manuel (favoriser en si présent)
    if not target_lang and manual_subs:
        target_lang = next((l for l in manual_subs if l.lower().startswith("en")), manual_subs[0])
        target_type = "manuel"

    # Priorité 4 : n'importe quel auto
    if not target_lang and include_auto and auto_subs:
        target_lang = next((l for l in auto_subs if l.lower().startswith("en")), auto_subs[0])
        target_type = "auto"

    if not target_lang:
        return None, None, "sous-titres auto désactivés par l'utilisateur"

    # Étape 3 : télécharger uniquement cette langue
    with tempfile.TemporaryDirectory() as tmp:
        dl_opts = {
            "quiet": True,
            "no_warnings": True,
            "skip_download": True,
            "writesubtitles": target_type == "manuel",
            "writeautomaticsub": target_type == "auto",
            "subtitleslangs": [target_lang],
            "subtitlesformat": "vtt",
            "outtmpl": os.path.join(tmp, "%(id)s.%(ext)s"),
        }
        try:
            with yt_dlp.YoutubeDL(dl_opts) as ydl:
                ydl.download([video_url])
        except Exception as e:
            return None, None, f"yt-dlp download : {type(e).__name__}"

        vtt_files = glob.glob(os.path.join(tmp, "*.vtt"))
        if not vtt_files:
            return None, None, f"fichier .vtt non créé pour la langue {target_lang}"

        text = parse_vtt(vtt_files[0])
        if not text:
            return None, None, "fichier .vtt vide après parsing"

        return text, f"{target_lang} ({target_type})", None


# ============================================================================
# FETCH - MÉTHODE 2 : youtube-transcript-api (fallback)
# ============================================================================
def fetch_transcript_api(video_id: str, lang_pref: list, include_auto: bool):
    """Fallback secondaire si yt-dlp n'a rien donné."""
    try:
        transcripts = YouTubeTranscriptApi.list_transcripts(video_id)
    except (TranscriptsDisabled, NoTranscriptFound):
        return None, None, "API : transcripts désactivés"
    except Exception as e:
        return None, None, f"API list : {type(e).__name__}"

    # Manuels en langue préférée
    for lang in lang_pref:
        try:
            t = transcripts.find_manually_created_transcript([lang])
            return " ".join(x["text"] for x in t.fetch()), f"{lang} (manuel-API)", None
        except Exception:
            continue

    # Auto en langue préférée
    if include_auto:
        for lang in lang_pref:
            try:
                t = transcripts.find_generated_transcript([lang])
                return " ".join(x["text"] for x in t.fetch()), f"{lang} (auto-API)", None
            except Exception:
                continue

    # Dernier recours : n'importe quoi de dispo
    try:
        for t in transcripts:
            if t.is_generated and not include_auto:
                continue
            try:
                text = " ".join(x["text"] for x in t.fetch())
                tag = "auto-API" if t.is_generated else "manuel-API"
                return text, f"{t.language_code} ({tag})", None
            except Exception:
                continue
    except Exception:
        pass

    return None, None, "API : aucun transcript récupérable"


# ============================================================================
# FETCH - ORCHESTRATEUR
# ============================================================================
def get_transcript(video_id: str, lang_pref: list, include_auto: bool):
    """Tente yt-dlp, puis API en fallback. Retourne (texte, langue, erreurs)."""
    errors = []

    text, lang, err = fetch_transcript_ytdlp(video_id, lang_pref, include_auto)
    if text:
        return text, lang, None
    if err:
        errors.append(f"yt-dlp → {err}")

    text, lang, err = fetch_transcript_api(video_id, lang_pref, include_auto)
    if text:
        return text, lang, None
    if err:
        errors.append(f"API → {err}")

    return None, None, " | ".join(errors) if errors else "raison inconnue"


# ============================================================================
# BUILD DOCX
# ============================================================================
def build_docx(source_title: str, channel_url: str, results: list) -> BytesIO:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(source_title, level=0)
    doc.add_paragraph(f"Source : {channel_url}")
    doc.add_paragraph(f"Nombre de vidéos : {len(results)}")
    doc.add_paragraph("")

    for idx, r in enumerate(results, 1):
        doc.add_heading(f"{idx}. {r['title']}", level=1)

        p = doc.add_paragraph()
        p.add_run("URL : ").bold = True
        p.add_run(r["url"])

        if r["transcript"]:
            p = doc.add_paragraph()
            p.add_run("Langue : ").bold = True
            p.add_run(r["lang"] or "inconnue")
            doc.add_paragraph(r["transcript"])
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"❌ Pas de transcription — {r.get('error', 'raison inconnue')}")
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def safe_filename(name: str) -> str:
    return re.sub(r"[^\w\-]+", "_", name)[:60].strip("_") or "transcripts"


# ============================================================================
# UI
# ============================================================================
col1, col2 = st.columns([3, 1])
with col1:
    channel_url = st.text_input(
        "URL de la chaîne ou de la playlist YouTube",
        placeholder="https://www.youtube.com/@NomDeLaChaine",
    )
with col2:
    num_videos = st.number_input("Nombre de vidéos", min_value=1, max_value=50, value=10)

col3, col4 = st.columns([2, 1])
with col3:
    lang_pref = st.multiselect(
        "Langues préférées (ordre prioritaire)",
        options=["fr", "en", "es", "de", "ar", "pt", "it", "ru", "zh", "ja"],
        default=["fr", "en"],
    )
with col4:
    include_auto = st.checkbox("Inclure sous-titres auto", value=True)

st.markdown("---")

if st.button("🚀 Extraire les transcriptions", type="primary", use_container_width=True):
    if not channel_url:
        st.error("Merci de fournir une URL")
        st.stop()
    if not lang_pref:
        st.error("Choisis au moins une langue")
        st.stop()

    # 1. Lister les vidéos
    with st.spinner(f"Récupération des {num_videos} dernières vidéos..."):
        url = normalize_channel_url(channel_url)
        try:
            videos, source_title = list_videos(url, int(num_videos))
        except Exception as e:
            st.error(f"Impossible de lister les vidéos : {e}")
            st.stop()

    if not videos:
        st.error("Aucune vidéo trouvée. Vérifie l'URL.")
        st.stop()

    st.success(f"✅ {len(videos)} vidéos trouvées dans **{source_title}**")

    # 2. Récupérer les transcriptions
    progress = st.progress(0.0)
    status = st.empty()
    results = []

    for i, v in enumerate(videos):
        status.info(f"📝 Transcription {i+1}/{len(videos)} — {v['title']}")
        text, lang, error = get_transcript(v["id"], lang_pref, include_auto)
        results.append({
            "id": v["id"],
            "title": v["title"],
            "url": v["url"],
            "transcript": text,
            "lang": lang,
            "error": error,
        })
        progress.progress((i + 1) / len(videos))

    status.empty()
    progress.empty()

    # 3. Métriques
    ok_count = sum(1 for r in results if r["transcript"])
    total_words = sum(len(r["transcript"].split()) for r in results if r["transcript"])

    c1, c2, c3 = st.columns(3)
    c1.metric("Transcriptions OK", f"{ok_count}/{len(results)}")
    c2.metric("Total mots", f"{total_words:,}".replace(",", " "))
    c3.metric("Échecs", len(results) - ok_count)

    # 4. Télécharger
    buf = build_docx(source_title, channel_url, results)
    filename = f"{safe_filename(source_title)}_transcripts.docx"
    st.download_button(
        "📥 Télécharger le .docx",
        data=buf,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True,
    )

    # 5. Détails des échecs (affichés direct, pas cachés)
    failures = [r for r in results if not r["transcript"]]
    if failures:
        with st.expander(f"⚠️ Détails des {len(failures)} échec(s)", expanded=True):
            for r in failures:
                st.markdown(f"- **[{r['title']}]({r['url']})** → `{r.get('error', 'n/a')}`")

    # 6. Aperçu
    with st.expander("👁️ Aperçu des transcriptions"):
        for r in results:
            st.markdown(f"### [{r['title']}]({r['url']})")
            if r["transcript"]:
                st.caption(f"Langue : {r['lang']} · {len(r['transcript'].split())} mots")
                preview = r["transcript"][:2000] + ("…" if len(r["transcript"]) > 2000 else "")
                st.text_area(" ", preview, height=150, key=r["id"], label_visibility="collapsed")
            else:
                st.warning(f"❌ {r.get('error', 'raison inconnue')}")
            st.markdown("---")
